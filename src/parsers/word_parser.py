"""
Word 文档解析器

支持提取文本、表格和段落结构
"""

from docx import Document
import pandas as pd
from pathlib import Path
from typing import Dict, Any, List
from .base_parser import BaseParser


class WordParser(BaseParser):
    """Word 文档解析器"""
    
    def __init__(self, file_path: str):
        """
        初始化 Word 解析器
        
        Args:
            file_path: Word 文件路径
        """
        super().__init__(file_path)
        self.doc = Document(str(self.file_path))
    
    def parse(self) -> Dict[str, Any]:
        """
        解析 Word 文档
        
        Returns:
            标准化数据结构
        """
        text = self.extract_text()
        tables = self.extract_tables()
        
        return {
            'metadata': self.get_metadata(),
            'content': {
                'text': text,
                'tables': tables,
                'structure': {
                    'paragraphs': len(self.doc.paragraphs),
                    'tables': len(self.doc.tables),
                    'sections': len(self.doc.sections)
                }
            },
            'metrics': {
                'word_count': len(text.split()),
                'character_count': len(text),
                'paragraph_count': len(self.doc.paragraphs),
                'table_count': len(self.doc.tables),
                'line_count': text.count('\n') + 1
            }
        }
    
    def extract_text(self) -> str:
        """
        提取所有文本内容
        
        Returns:
            文档的纯文本
        """
        paragraphs = []
        
        for para in self.doc.paragraphs:
            text = para.text.strip()
            if text:  # 忽略空段落
                paragraphs.append(text)
        
        return '\n'.join(paragraphs)
    
    def extract_tables(self) -> List[pd.DataFrame]:
        """
        提取所有表格
        
        Returns:
            表格列表（DataFrame 格式）
        """
        tables = []
        
        for table_idx, table in enumerate(self.doc.tables):
            # 提取表格数据
            data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                data.append(row_data)
            
            if not data or len(data) < 2:
                continue  # 跳过空表格或只有一行的表格
            
            try:
                # 第一行作为表头
                headers = data[0]
                rows = data[1:]
                
                # 创建 DataFrame
                df = pd.DataFrame(rows, columns=headers)
                
                # 添加表格索引信息
                df.attrs['table_index'] = table_idx
                df.attrs['source'] = 'word_document'
                
                tables.append(df)
                
            except Exception as e:
                print(f"  ⚠️  表格 {table_idx + 1} 解析失败: {e}")
                continue
        
        return tables
    
    def extract_headings(self) -> List[Dict[str, Any]]:
        """
        提取标题结构
        
        Returns:
            标题列表，包含级别和内容
        """
        headings = []
        
        for para in self.doc.paragraphs:
            if para.style.name.startswith('Heading'):
                level = int(para.style.name.replace('Heading ', ''))
                headings.append({
                    'level': level,
                    'text': para.text.strip()
                })
        
        return headings


def main():
    """测试函数"""
    import sys
    
    if len(sys.argv) < 2:
        print("用法: python word_parser.py <docx文件路径>")
        return
    
    file_path = sys.argv[1]
    
    parser = WordParser(file_path)
    data = parser.parse()
    
    print(f"\n✅ 解析成功:")
    print(f"  - 文件: {data['metadata']['file_name']}")
    print(f"  - 大小: {data['metadata']['file_size_mb']} MB")
    print(f"  - 段落数: {data['metrics']['paragraph_count']}")
    print(f"  - 表格数: {data['metrics']['table_count']}")
    print(f"  - 字数: {data['metrics']['word_count']}")
    
    # 显示前 500 字符
    text = data['content']['text']
    if len(text) > 500:
        print(f"\n📝 内容预览:\n{text[:500]}...")
    else:
        print(f"\n📝 内容:\n{text}")


if __name__ == '__main__':
    main()
